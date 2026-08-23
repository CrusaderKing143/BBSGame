from __future__ import annotations

import sys
from pathlib import Path

from docx import Document
from docx.text.paragraph import Paragraph


OUTPUT_PATH = Path("Docs/BBSGame 全流程旁白脚本_黄色词吐槽补充版.docx")


INSERTIONS = [
    (
        "中文\n连鞋底都量了。\n我低头看了一眼自己的条纹袜子，突然有点没有安全感。",
        [
            (
                "L1-C-01｜点击黄色词 “platform shoes”",
                "Platform shoes. Great. We’re measuring guilt in centimeters now.",
                "厚底鞋。很好，现在嫌疑已经开始按厘米算了。",
            ),
            (
                "L1-C-02｜点击黄色词 “bag”",
                "Same bag, same charm. Apparently accessories can testify now.",
                "同款包，同一个挂饰。看来配饰也能出庭作证了。",
            ),
        ],
    ),
    (
        "中文\n私人账号关注了夜店。照片里的人手腕上也缠着东西。",
        [
            (
                "L1-C-03｜点击黄色词 “manager”",
                "His manager. Or someone with the same back and a very unlucky haircut.",
                "他的经纪人。或者只是后脑勺和发型都很倒霉的路人。",
            ),
            (
                "L1-C-04｜点击黄色词 “bar operator”",
                "The club’s official account. One follow, and suddenly we have a relationship chart.",
                "夜店官方账号。点一下关注，关系图就自动连上了。",
            ),
            (
                "L1-C-05｜点击黄色词 “account”",
                "A private account. Privacy lasts right up until somebody needs a clue.",
                "私人账号。隐私通常只维持到有人缺线索为止。",
            ),
            (
                "L1-C-06｜点击黄色词 “wrist”",
                "A wrist. We’re investigating body parts now. Very normal.",
                "一截手腕。我们已经开始拿身体部位办案了。很正常。",
            ),
            (
                "L1-C-07｜点击黄色词 “brace”",
                "A wrist brace. Could be REN’s. Could be from every pharmacy in town.",
                "一只护腕。可能是REN的，也可能全城药店都有卖。",
            ),
        ],
    ),
    (
        "中文\n录音、排练、第二天一早的航班……\n这行程排得比我的午休还挤。",
        [
            (
                "L1-C-08｜点击黄色词 “handmade doll”",
                "Handmade doll. Cute. Also, apparently, an alibi with ears.",
                "手工玩偶。挺可爱，顺便还是个长耳朵的不在场证明。",
            ),
            (
                "L1-C-09｜点击黄色词 “lyric notebook”",
                "A lyric notebook. Nothing says innocence like working after midnight.",
                "歌词本。凌晨还在工作，确实很有无辜的说服力。",
            ),
            (
                "L1-C-10｜点击黄色词 “coffee”",
                "Coffee. Finally, a witness I understand.",
                "咖啡。终于来了一个我能理解的证人。",
            ),
            (
                "L1-C-11｜点击黄色词 “schedule”",
                "A packed schedule. Mine looks similar, minus the airports and dignity.",
                "塞满的行程表。我的也差不多，只是少了机场和体面。",
            ),
        ],
    ),
    (
        "中文\n而REN的名字，正安安静静躺在公司的股权记录里。",
        [
            (
                "L2-C-01｜点击黄色词 “old interview”",
                "An old interview. The internet’s favorite time machine; it only stops where convenient.",
                "旧采访。互联网最爱的时光机，只停在方便的地方。",
            ),
            (
                "L2-C-02｜点击黄色词 “REN”",
                "REN’s name, highlighted. Subtle. Practically a verdict.",
                "REN的名字还特地标黄。真含蓄，差不多就差宣判了。",
            ),
            (
                "L2-C-03｜点击黄色词 “The VIBE”",
                "The VIBE. A club name that sounds like evidence before I even click it.",
                "The VIBE。夜店名字听着就像证物，挺省文案。",
            ),
            (
                "L2-C-04｜点击黄色词 “watch box”",
                "A watch box. Rich people’s clutter photographs remarkably well.",
                "一个表盒。有钱人的杂物，上镜都很像线索。",
            ),
        ],
    ),
    (
        "中文\n这家夜店甚至提前知道了所谓的突击检查。",
        [
            (
                "L2-C-05｜点击黄色词 “surveillance hard drive”",
                "Surveillance hard drive. Broken, obviously. A working camera would ruin the genre.",
                "监控硬盘。坏了，当然。监控要是好好的，这故事就不成立了。",
            ),
            (
                "L2-C-06｜点击黄色词 “shift schedule”",
                "Shift schedule. Add enough arrows, and even clocking in looks criminal.",
                "排班表。箭头画得够多，连打卡都能显得可疑。",
            ),
        ],
    ),
    (
        "中文\n他的官方行程从早排到晚。\n光是看着，我都替他累。",
        [
            (
                "L2-C-07｜点击黄色词 “resignation letter”",
                "A resignation letter dated early. One little date, and the accusation starts wobbling.",
                "一封更早的辞职信。一个日期，就把整套指控晃松了。",
            ),
            (
                "L2-C-08｜点击黄色词 “official schedule”",
                "Official schedule. Very neat. Official things do love neatness.",
                "官方行程。整整齐齐。官方的东西就是很爱整齐。",
            ),
        ],
    ),
    (
        "中文\n伤者可能先闹了事。REN留在夜店的东西也早就被收拾退回。",
        [
            (
                "L2-C-09｜点击黄色词 “news report”",
                "A news report says the victim swung first. There goes the clean version.",
                "新闻说是伤者先动手。好，干净利落的版本没了。",
            ),
            (
                "L2-C-10｜点击黄色词 “limited-edition collab toy”",
                "Limited-edition collab toy. Even the alibi has merchandise.",
                "限定联名玩偶。连不在场证明都有周边了。",
            ),
            (
                "L2-C-11｜点击黄色词 “car keys”",
                "Car keys returned before the incident. Tiny object, enormous narrative responsibility.",
                "事发前就退回的车钥匙。小小一件东西，扛着好大的叙事责任。",
            ),
        ],
    ),
    (
        "中文\n车里的一顶帽子，桌上的一枚戒指。\n互联网侦探重新上岗了。",
        [
            (
                "L3-C-01｜点击黄色词 “baseball cap”",
                "A baseball cap. Celebrity investigations remain bravely committed to hats.",
                "棒球帽。娱乐圈侦查始终坚定地相信帽子。",
            ),
            (
                "L3-C-02｜点击黄色词 “custom ring”",
                "A custom ring. At least this clue had the decency to look expensive.",
                "定制戒指。至少这条线索很懂礼貌，贵得很明显。",
            ),
        ],
    ),
    (
        "中文\n有钱人留下的线索，看起来也格外昂贵。",
        [
            (
                "L3-C-03｜点击黄色词 “paparazzi”",
                "A paparazzo. The only person here whose job makes mine look stable.",
                "狗仔。这里唯一一个能让我的工作显得稳定的人。",
            ),
            (
                "L3-C-04｜点击黄色词 “confidentiality agreement”",
                "Confidentiality agreement. Nothing says ‘nothing happened’ like legal paperwork.",
                "保密协议。最能证明“什么都没发生”的，当然是法律文件。",
            ),
            (
                "L3-C-05｜点击黄色词 “check”",
                "That check has more zeroes than my annual review.",
                "这张签单上的零，比我年度考核里好看的数字还多。",
            ),
        ],
    ),
    (
        "中文\n一份商业合同。看来REN和Leo确实合作过。",
        [
            (
                "L3-C-06｜点击黄色词 “interview”",
                "An interview. Leo says business. The comments will hear confession.",
                "一段采访。Leo说的是生意，评论区听见的会是口供。",
            ),
            (
                "L3-C-07｜点击黄色词 “business contract”",
                "A business contract. Boring, stamped, and annoyingly reasonable.",
                "商业合同。无聊、盖章，而且合理得让人烦。",
            ),
        ],
    ),
    (
        "中文\n助理说他当晚病了。药物和视频都能证明这一点。",
        [
            (
                "L3-C-08｜点击黄色词 “assistant”",
                "His assistant. A human alibi in office clothes.",
                "他的助理。一个穿着职业装的人形不在场证明。",
            ),
            (
                "L3-C-09｜点击黄色词 “hospital wristband”",
                "Hospital wristband. Harder to argue with than a blurry cap.",
                "医院手环。这个比一顶模糊的帽子难反驳多了。",
            ),
            (
                "L3-C-10｜点击黄色词 “medicine”",
                "Medicine. The least glamorous item in the folder, so probably the most useful.",
                "药。文件夹里最不上镜的东西，偏偏可能最有用。",
            ),
        ],
    ),
]

SELECTED_HEADINGS = {
    "L1-C-02｜点击黄色词 “bag”",
    "L1-C-03｜点击黄色词 “manager”",
    "L1-C-10｜点击黄色词 “coffee”",
    "L2-C-01｜点击黄色词 “old interview”",
    "L2-C-05｜点击黄色词 “surveillance hard drive”",
    "L2-C-07｜点击黄色词 “resignation letter”",
    "L3-C-01｜点击黄色词 “baseball cap”",
    "L3-C-04｜点击黄色词 “confidentiality agreement”",
    "L3-C-07｜点击黄色词 “business contract”",
}


def insert_after(paragraph: Paragraph, text: str, style_name: str) -> Paragraph:
    new_paragraph = paragraph._parent.add_paragraph(text, style=style_name)
    paragraph._p.addnext(new_paragraph._p)
    return new_paragraph


def find_unique_paragraph(document: Document, text: str) -> Paragraph:
    matches = [paragraph for paragraph in document.paragraphs if paragraph.text == text]
    if len(matches) != 1:
        raise RuntimeError(f"Expected one anchor paragraph, found {len(matches)}: {text!r}")
    return matches[0]


def add_bilingual_trigger(
    anchor: Paragraph,
    heading: str,
    english: str,
    chinese: str,
) -> Paragraph:
    current = insert_after(anchor, heading, "Heading 3")
    current.paragraph_format.keep_with_next = True
    current = insert_after(current, f"EN\n{english}", "Normal (Web)")
    current.paragraph_format.keep_with_next = True
    current = insert_after(current, f"中文\n{chinese}", "Normal (Web)")
    current.paragraph_format.keep_together = True
    return current


def main() -> None:
    if len(sys.argv) != 2:
        raise SystemExit("Usage: extend_yellow_word_narration.py <source.docx>")

    source_path = Path(sys.argv[1])
    document = Document(source_path)

    for anchor_text, triggers in INSERTIONS:
        triggers = [trigger for trigger in triggers if trigger[0] in SELECTED_HEADINGS]
        if not triggers:
            continue
        current = find_unique_paragraph(document, anchor_text)
        for heading, english, chinese in triggers:
            current = add_bilingual_trigger(current, heading, english, chinese)

    # The source ends with an empty paragraph. After the grouped bilingual
    # trigger reflows the document, Word can push that empty paragraph onto a
    # completely blank final page, so omit this non-content paragraph.
    if document.paragraphs and document.paragraphs[-1].text == "":
        trailing_paragraph = document.paragraphs[-1]._element
        trailing_paragraph.getparent().remove(trailing_paragraph)

    OUTPUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    document.save(OUTPUT_PATH)

    verified = Document(OUTPUT_PATH)
    output_texts = [paragraph.text for paragraph in verified.paragraphs]
    source_texts = [paragraph.text for paragraph in Document(source_path).paragraphs]
    if source_texts and source_texts[-1] == "":
        source_texts.pop()
    cursor = 0
    for source_text in source_texts:
        while cursor < len(output_texts) and output_texts[cursor] != source_text:
            cursor += 1
        if cursor >= len(output_texts):
            raise RuntimeError(f"Original paragraph was not preserved: {source_text!r}")
        cursor += 1

    inserted_headings = [text for text in output_texts if text in SELECTED_HEADINGS]
    if len(inserted_headings) != len(SELECTED_HEADINGS):
        raise RuntimeError(
            f"Expected {len(SELECTED_HEADINGS)} click triggers, found {len(inserted_headings)}."
        )

    print(
        f"Verified {len(source_texts)} original paragraphs and "
        f"{len(inserted_headings)} inserted click triggers."
    )
    print(OUTPUT_PATH.resolve())


if __name__ == "__main__":
    main()
