from __future__ import annotations

import json
import sys
import argparse
from pathlib import Path

from docx import Document


def run_data(run):
    color = run.font.color.rgb
    return {
        "text": run.text,
        "bold": run.bold,
        "italic": run.italic,
        "underline": run.underline,
        "highlight": str(run.font.highlight_color) if run.font.highlight_color else None,
        "color": str(color) if color else None,
        "font": run.font.name,
        "size_pt": run.font.size.pt if run.font.size else None,
    }


def paragraph_data(paragraph, index=None):
    return {
        "index": index,
        "style": paragraph.style.name if paragraph.style else None,
        "text": paragraph.text,
        "runs": [run_data(run) for run in paragraph.runs],
    }


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("source")
    parser.add_argument("--start", type=int, default=0)
    parser.add_argument("--end", type=int)
    parser.add_argument("--compact", action="store_true")
    args = parser.parse_args()

    source = Path(args.source)
    document = Document(source)

    if args.compact:
        end = args.end if args.end is not None else len(document.paragraphs)
        for index in range(args.start, min(end, len(document.paragraphs))):
            paragraph = document.paragraphs[index]
            text = paragraph.text.replace("\t", " ").replace("\n", " / ")
            highlights = [
                f"{run.text}:{run.font.highlight_color}"
                for run in paragraph.runs
                if run.font.highlight_color
            ]
            print(
                f"{index:03d}\t{paragraph.style.name if paragraph.style else ''}\t{text}"
                + (f"\tHIGHLIGHT={' | '.join(highlights)}" if highlights else "")
            )
        return

    payload = {
        "source": str(source),
        "paragraphs": [
            paragraph_data(paragraph, index)
            for index, paragraph in enumerate(document.paragraphs)
        ],
        "tables": [],
        "sections": [],
    }

    for table_index, table in enumerate(document.tables):
        table_data = {"index": table_index, "rows": []}
        for row_index, row in enumerate(table.rows):
            row_data = {"index": row_index, "cells": []}
            for cell_index, cell in enumerate(row.cells):
                row_data["cells"].append(
                    {
                        "index": cell_index,
                        "paragraphs": [paragraph_data(p) for p in cell.paragraphs],
                    }
                )
            table_data["rows"].append(row_data)
        payload["tables"].append(table_data)

    for index, section in enumerate(document.sections):
        payload["sections"].append(
            {
                "index": index,
                "width_in": section.page_width.inches,
                "height_in": section.page_height.inches,
                "top_margin_in": section.top_margin.inches,
                "bottom_margin_in": section.bottom_margin.inches,
                "left_margin_in": section.left_margin.inches,
                "right_margin_in": section.right_margin.inches,
            }
        )

    print(json.dumps(payload, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
